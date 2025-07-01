from docx import Document
import math2docx

def write_docx_from_structured_text(structured_text, output_path):
    """
    structured_text: List of dicts with 'text' and 'is_equation' keys.
    output_path: Path to save the DOCX file.
    """
    document = Document()
    for item in structured_text:
        if item.get('is_equation'):
            paragraph = document.add_paragraph()
            math2docx.add_math(paragraph, item['text'])
        else:
            document.add_paragraph(item['text'])
    document.save(output_path)

def write_latex_to_docx(latex_code, output_path):
    """
    Inserts the full LaTeX code as a single math block into the DOCX file.
    """
    document = Document()
    paragraph = document.add_paragraph()
    math2docx.add_math(paragraph, latex_code)
    document.save(output_path)

def write_explanation_and_equation_to_docx(explanation, equation, output_path):
    """
    Writes the explanation as normal text and the equation as a math block to the DOCX file.
    """
    document = Document()
    if explanation:
        document.add_paragraph(explanation)
    if equation:
        paragraph = document.add_paragraph()
        math2docx.add_math(paragraph, equation)
    document.save(output_path) 